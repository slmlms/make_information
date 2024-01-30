import os
import docx

def rename_documents(directory):
  """
  Renames all Word documents in a directory to the first line of the document.

  Args:
    directory: The directory containing the Word documents.
  """
  for filename in os.listdir(directory):
    filepath = os.path.join(directory, filename)
    if not os.path.isfile(filepath) or not filepath.endswith(".docx"):
      continue

    # Open the document and read the first line
    try:
      doc = docx.Document(filepath)
      first_line = doc.paragraphs[0].text.strip()
    except Exception as e:
      print(f"Error opening document '{filepath}': {e}")
      continue

    # Remove illegal characters from the filename
    new_filename = "".join(ch for ch in first_line if ch.isalnum() or ch == " ")

    # Build the new filepath
    new_filepath = os.path.join(directory, f"{new_filename}.docx")

    # Rename the file
    try:
      os.rename(filepath, new_filepath)
      print(f"Renamed '{filepath}' to '{new_filepath}'")
    except Exception as e:
      print(f"Error renaming file '{filepath}': {e}")

# Replace "your_directory" with the actual directory path
# rename_documents("D:\Jobs\卡莫亚\检验批及分项\模板\钢结构安装\检验批")



import os
from docx import Document



def find_actual_next_cell(row, current_cell):
    """
    Find the actual next cell considering merged cells.
    Returns None if there is no next non-merged cell.
    """
    # 假设cells是按顺序排列的，且当前cell已经在循环中被正确识别
    current_col_idx = row.cells.index(current_cell)
    total_cols = len(row.cells)

    for i in range(current_col_idx + 1, total_cols):
        cell = row.cells[i]
        colspan = 1  # 默认单个单元格跨度为1

        # 根据单元格的边界计算实际跨度
        # 注意：这里假设了合并单元格的边界计算方法，实际情况请根据docx库的具体行为调整
        if hasattr(cell, 'merge'):
            # 如果有merge属性，则尝试获取合并范围
            # 这里假设merge返回一个包含起始和结束列索引的元组
            # 实际上docx并没有直接提供这样的属性，此处仅作为演示如何处理合并单元格
            merge_range = cell.merge  # 示例代码，实际应根据库的功能进行调整
            colspan = merge_range[1] - merge_range[0] + 1

        # 如果当前单元格的跨度已经覆盖到下一列，则跳过它
        if colspan > 1 and (current_col_idx + 1) < (i + colspan):
            continue
        else:
            return cell
    return None
def process_documents(directory):
    """
    Iterates through Word documents in a directory, finds and modifies text as specified.

    Args:
        directory: The directory containing the Word documents.
    """
    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        if not os.path.isfile(filepath) or not filepath.endswith(".docx"):
            continue

        try:
            doc = Document(filepath)

            # Find and modify "编号："
            for paragraph in doc.paragraphs:
                if "编号：" in paragraph.text:
                    paragraph.text += " {{Bianhao}}"

            # Find and modify table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text == "单位工程名称":
                            next_cell = find_actual_next_cell(row, cell)
                            if next_cell:
                                next_cell.text = "{{Zixiangmingcheng}}"
                        elif cell.text == "检验批容量":
                            next_cell = find_actual_next_cell(row, cell)
                            if next_cell:
                                next_cell.text = "{{Jianyanpirongliang}}"
                        elif cell.text == "检验批部位":
                            next_cell = find_actual_next_cell(row, cell)
                            if next_cell:
                                next_cell.text = "{{Jianyanpibuwei}}"
            doc.save(filepath)
            print(f"Processed '{filepath}'")

        except Exception as e:
            print(f"Error processing document '{filepath}': {e}")

# Replace "your_folder_path" with the actual folder path
process_documents("D:\Jobs\卡莫亚\检验批及分项\模板\钢结构安装\检验批\测试")

