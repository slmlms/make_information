import random
from docx import Document

# 已有方法生成随机整数并转换为字符串
def generate_random_integer_and_convert_to_string(start, end):
    random_int = random.randint(start, end)
    return str(random_int)

# 新建函数以生成指定长度的随机数数组
def generate_random_number_array(length, start=1, end=9):
    random_numbers = [generate_random_integer_and_convert_to_string(start, end) for _ in range(length)]
    return random_numbers

# 修改fill_table_with_random_numbers函数以使用预先生成的数组
def fill_table_with_predefined_numbers(doc_path, table_index, row_range, col_range, number_array):
    # 打开Word文档
    doc = Document(doc_path)

    # 获取指定索引的表格
    tables = doc.tables
    if table_index < len(tables):
        table = tables[table_index]
        array_index = 0

        for r in range(*row_range):
            print(r)
            row_cells = table.rows[r].cells
            for c in range(*col_range):
                print(c)
                # 使用数组中的下一个随机数填入表格单元格
                if array_index < len(number_array):
                    random_str = number_array[array_index]
                    cell = row_cells[c]
                    p = cell.add_paragraph()
                    p.add_run(random_str)
                    array_index += 1
                else:
                    break  # 如果数组用尽，则停止填充

        # 保存更改后的文档
        doc.save(doc_path)
    else:
        print(f"Table index {table_index} is out of the document's table count.")

# 使用示例：先生成长度为10的随机数数组，再填充到表格中
random_number_list = generate_random_number_array(10)
fill_table_with_predefined_numbers("D:\\Jobs\\卡莫亚\\检验批及分项\\钢结构检验批生成\\精矿脱水车间\\钢结构安装\\钢结构压型钢板\\钢结构压型钢板检验批质量检验评定记录-屋面瓦-验收部位-屋面瓦.docx",
                                  0, (13,14), (6,16), random_number_list)
