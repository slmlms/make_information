import os

import tqdm
from docx import Document
from docx.shared import Pt

def replace_text_in_docx(directory, old_text, new_text):
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(('.docx', '.DOCX')):
                full_path = os.path.join(root, file)
                # 检查文件名是否包含指定字符串
                if "分项质量检查验收记录" in file:
                    print(f"正在处理文件：{full_path}")

                    # 打开并修改docx文档
                    doc = Document(full_path)
                    # 遍历文档中的所有表格
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                # 遍历单元格内的所有段落
                                for paragraph in cell.paragraphs:
                                    if old_text in paragraph.text:
                                        paragraph.text = paragraph.text.replace(old_text, new_text)
                                        # 设置字号为小五
                                        runs_after_replace = paragraph.runs
                                        if runs_after_replace:
                                            # 设置最后一个运行的字体大小
                                            runs_after_replace[-1].font.size = Pt(9)

                    # 保存更改后的文档
                    doc.save(full_path)


# 使用方法
directory_path = 'D:\Jobs\卡莫亚\检验批及分项\二期选矿竣工资料'  # 替换为你的目录路径
old_string = "EPC总承包单位"
new_string = "总承包单位"

tqdm.tqdm(replace_text_in_docx(directory_path, old_string, new_string))
