import os
import subprocess

from docx import Document


def convert_to_docx(input_dir):
    # 遍历指定目录下的Word 97-2003文档（.doc）
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith(".doc"):
                doc_path = os.path.join(root, file)
                # 构建输出路径
                docx_path = os.path.splitext(doc_path)[0] + ".docx"
                # 使用Libreoffice将.doc文件转换为.docx文件
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", input_dir])
                # 删除原始的.doc文件
                os.remove(doc_path)


def rename_docx_files(input_dir):
    # 遍历转换完成的.docx文件
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith(".docx"):
                docx_path = os.path.join(root, file)
                # 读取文档首行内容
                doc = Document(docx_path)
                if doc.paragraphs:
                    first_line = doc.paragraphs[0].text.strip()
                else:
                    first_line = ""

                # 构建新的文件名
                new_name = os.path.join(root, first_line + ".docx")
                index = 1
                while os.path.exists(new_name):
                    new_name = os.path.join(root, first_line + "-" + str(index) + ".docx")
                    index += 1

                # 重命名.docx文件
                os.rename(docx_path, new_name)


# 指定目录路径
input_directory = "D:\Jobs\图集规范\建筑工程资料(下载到电脑上解压)\地方版资料（陆续更新）\湖南建筑资料员表格\湖南省建筑工程全套资料表格\湖南省建筑工程全套资料表格\检验批验收表\\08(39)"
# 转换为.docx格式
# convert_to_docx(input_directory)
# 将文档首行内容作为文件名保存
rename_docx_files(input_directory)
